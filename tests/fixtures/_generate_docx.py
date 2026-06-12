"""Generate a tiny DOCX fixture. Run once; commit the output."""
from __future__ import annotations

from pathlib import Path

from docx import Document

out = Path(__file__).parent / "tiny.docx"
doc = Document()
doc.add_paragraph("First paragraph for viewer tests.")
doc.add_paragraph("Second paragraph with a bit more text.")
doc.add_paragraph("Third paragraph closes the tiny document.")
doc.save(out)
print(f"wrote {out} ({out.stat().st_size} bytes)")
